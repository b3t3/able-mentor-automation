import os
import csv
import docx
import math


CURRENT_DIRECTORY = os.path.dirname(os.path.realpath(__file__)).replace("\\", "/")
OUTPUT_DIRECTORY = f"{CURRENT_DIRECTORY}/student_profiles"
REGISTER_FILE_NAME = "student_register.csv"
REGISTER_FILE_PATH = f"{CURRENT_DIRECTORY}/{REGISTER_FILE_NAME}"


def get_column_index(column):
    # 26 number system where [A...Z] is mapped to [1...26]
    decimal_value = 0
    for idx in reversed(range(0, len(column))):
        decimal_value += (ord(column[idx]) - ord("A") + 1) * math.pow(26, len(column) - idx - 1)

    return int(decimal_value - 1)  # the index is the decimal value minus 1


CONFIRMED = get_column_index("A")
STUDENT_NAME = get_column_index("AB")
AGE = get_column_index("AC")
CITY = get_column_index("AG")
SCHOOL_NAME = get_column_index("AK")
GRADE = get_column_index("AL")
SCHOOL_INTERESTS = get_column_index("AM")
NON_SCHOOL_INTERESTS = get_column_index("AO")
MOTIVATION = get_column_index("AS")
SIMILAR_PROGRAMS = get_column_index("AT")
NEED = get_column_index("AX")
WHAT_TO_DO_AFTER_SCHOOL = get_column_index("AN")
INTERESTS = get_column_index("AV")
DIFFICULT_SITUATION = get_column_index("AP")
WHY_APPLY = get_column_index("AW")
HOURS_PER_WEEK = get_column_index("AR")
PROJECT_SPHERE = get_column_index("AU")
HEARD_OF_ABLE_MENTOR = get_column_index("BB")
ROLE_IN_TEAM = get_column_index("AQ")
WANT_TO_IMPROVE = get_column_index("AY")
SPECIFIC_SPHERE = get_column_index("AZ")
ASPECTS_FOR_WORKING_WITH_MENTOR = get_column_index("BA")

column_titles = {
    CONFIRMED: "гр",
    STUDENT_NAME: "Ученик",
    AGE: "Възраст",
    CITY: "Населено място",
    SCHOOL_NAME: "Училище",
    GRADE: "Завършен клас",
    SCHOOL_INTERESTS: "Любими предмети",
    WHAT_TO_DO_AFTER_SCHOOL: "Планове след гимназията:",
    NON_SCHOOL_INTERESTS: "Интереси извън училище",
    DIFFICULT_SITUATION: "Лично предизвикателство, което си преодолял/а",
    ROLE_IN_TEAM: "Роля в екип",
    HOURS_PER_WEEK: "Часове седмично, които можеш да отделиш за програмата",
    MOTIVATION: "Мотивация за участие в програмата",
    SIMILAR_PROGRAMS: "Участвал ли си в други сходни програми, завършил ли си ги и какво си взе от тях?", #F
    PROJECT_SPHERE: "Проектни сфери",
    INTERESTS: "Кариерни и академичи насоки",
    WHY_APPLY: "Защо кандидатстваш в програмата?",
    NEED: "Уточнение към цел на участие(ако има)?",
    WANT_TO_IMPROVE: "Сфери за развитие",
    SPECIFIC_SPHERE: "Специфична сфера",
    ASPECTS_FOR_WORKING_WITH_MENTOR: "Аспекти за работа с ментор",
    HEARD_OF_ABLE_MENTOR: "Научил/а за ABLE Mentor от?",
}

def try_create_doc(row_data, file_path):
    if row_data[CONFIRMED] != "matched":
        return False

    doc = docx.Document()
    doc.add_heading("ПРОФИЛ НА ТВОЯ УЧЕНИК", 0)

    table = doc.add_table(rows=0, cols=2)

    for idx in range(0, len(row_data)):
        if (idx == CONFIRMED or
            idx == STUDENT_NAME or
                idx not in column_titles):
            continue

        if idx > HEARD_OF_ABLE_MENTOR:
            break

        table_row = table.add_row().cells
        table_row[0].text = column_titles[idx]
        table_row[1].text = row_data[idx]

    doc.save(file_path)
    return True


def create_docs():
    if not os.path.exists(OUTPUT_DIRECTORY):
        os.mkdir(OUTPUT_DIRECTORY)

    doc_counter = 1  # Start a counter for numbering the files

    with open(REGISTER_FILE_PATH, encoding="utf-8", mode="r") as fstream:
        reader = csv.reader(fstream, delimiter=',', quotechar='"')

        for idx, row in enumerate(reader):
            if idx == 0:
                continue

            student_name = row[STUDENT_NAME]
            file_path = f"{OUTPUT_DIRECTORY}/{doc_counter}_{student_name}.docx"

            if try_create_doc(row, file_path):
                doc_counter += 1  # Increment the counter only if the document is created successfully


if __name__ == "__main__":
    create_docs()
